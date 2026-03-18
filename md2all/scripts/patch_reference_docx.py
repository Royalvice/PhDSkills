#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path

from common import ASSETS_DIR, load_json


def load_profile(template_id: str) -> dict:
    catalog = load_json(ASSETS_DIR / "reference-docx" / "template-catalog.json")
    for item in catalog["templates"]:
        if item["id"] == template_id:
            return load_json(ASSETS_DIR / "reference-docx" / item["profile"])
    raise KeyError(f"Unknown template: {template_id}")


def patch_profile(profile: dict, overrides: dict[str, str]) -> dict:
    patched = json.loads(json.dumps(profile))
    if overrides.get("font"):
        patched["body"]["font_zh"] = overrides["font"]
        patched["body"]["font_en"] = overrides["font"]
    if overrides.get("font_size"):
        patched["body"]["size_pt"] = float(overrides["font_size"])
    if overrides.get("margin_cm"):
        patched["page"]["margin_cm"] = float(overrides["margin_cm"])
    return patched


def generate_docx(profile: dict, output: Path) -> None:
    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise RuntimeError("python-docx is required to generate reference.docx templates") from exc

    def ensure_child(parent, tag: str):
        child = parent.find(qn(tag))
        if child is None:
            child = OxmlElement(tag)
            parent.append(child)
        return child

    def set_style_fonts(style, font_zh: str, font_en: str) -> None:
        style.font.name = font_en
        r_pr = style.element.get_or_add_rPr()
        r_fonts = ensure_child(r_pr, "w:rFonts")
        for key in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
            if qn(key) in r_fonts.attrib:
                del r_fonts.attrib[qn(key)]
        r_fonts.set(qn("w:ascii"), font_en)
        r_fonts.set(qn("w:hAnsi"), font_en)
        r_fonts.set(qn("w:cs"), font_en)
        r_fonts.set(qn("w:eastAsia"), font_zh)

    def set_spacing(style, line_spacing: float, space_before_pt: float, space_after_pt: float, align: str | None = None) -> None:
        style.paragraph_format.line_spacing = line_spacing
        style.paragraph_format.space_before = Pt(space_before_pt)
        style.paragraph_format.space_after = Pt(space_after_pt)
        if align == "center":
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "left":
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == "justify":
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif align == "right":
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def ensure_paragraph_style(document, name: str, base_name: str | None = None):
        try:
            style = document.styles[name]
        except KeyError:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        if base_name:
            try:
                style.base_style = document.styles[base_name]
            except KeyError:
                pass
        return style

    def configure_body_style(style, font_zh: str, font_en: str, size_pt: float, line_spacing: float, space_before_pt: float,
                             space_after_pt: float, first_line_indent_pt: float, align: str) -> None:
        set_style_fonts(style, font_zh, font_en)
        style.font.size = Pt(size_pt)
        set_spacing(style, line_spacing, space_before_pt, space_after_pt, align=align)
        style.paragraph_format.first_line_indent = Pt(first_line_indent_pt)

    def configure_heading_style(style, font_zh: str, font_en: str, size_pt: float, bold: bool, space_before_pt: float,
                                space_after_pt: float, line_spacing: float, align: str) -> None:
        set_style_fonts(style, font_zh, font_en)
        style.font.size = Pt(size_pt)
        style.font.bold = bold
        style.font.italic = False
        set_spacing(style, line_spacing, space_before_pt, space_after_pt, align=align)
        r_pr = style.element.get_or_add_rPr()
        color = ensure_child(r_pr, "w:color")
        color.set(qn("w:val"), "000000")
        for key in ("w:themeColor", "w:themeTint", "w:themeShade"):
            if qn(key) in color.attrib:
                del color.attrib[qn(key)]

    def configure_caption_style(style, font_zh: str, font_en: str, size_pt: float, line_spacing: float,
                                space_before_pt: float, space_after_pt: float, align: str) -> None:
        set_style_fonts(style, font_zh, font_en)
        style.font.size = Pt(size_pt)
        style.font.bold = False
        style.font.italic = False
        r_pr = style.element.get_or_add_rPr()
        for tag in ("w:i", "w:iCs"):
            node = r_pr.find(qn(tag))
            if node is not None:
                r_pr.remove(node)
        set_spacing(style, line_spacing, space_before_pt, space_after_pt, align=align)
        style.paragraph_format.first_line_indent = Pt(0)

    def configure_title_style(style, font_zh: str, font_en: str) -> None:
        set_style_fonts(style, font_zh, font_en)
        style.font.size = Pt(26)
        style.font.bold = False
        style.font.italic = False
        set_spacing(style, 1.15, 0, 12, align="left")
        style.paragraph_format.first_line_indent = Pt(0)
        r_pr = style.element.get_or_add_rPr()
        for tag in ("w:i", "w:iCs"):
            node = r_pr.find(qn(tag))
            if node is not None:
                r_pr.remove(node)
        color = ensure_child(r_pr, "w:color")
        color.set(qn("w:val"), "000000")
        for key in ("w:themeColor", "w:themeTint", "w:themeShade"):
            if qn(key) in color.attrib:
                del color.attrib[qn(key)]
        p_pr = style.element.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is not None:
            p_pr.remove(p_bdr)

    def configure_doc_defaults(document, font_zh: str, font_en: str, size_pt: float, line_spacing: float,
                               space_before_pt: float, space_after_pt: float) -> None:
        styles_element = document.styles.element
        doc_defaults = ensure_child(styles_element, "w:docDefaults")
        r_pr_default = ensure_child(doc_defaults, "w:rPrDefault")
        r_pr = ensure_child(r_pr_default, "w:rPr")
        r_fonts = ensure_child(r_pr, "w:rFonts")
        for key in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
            if qn(key) in r_fonts.attrib:
                del r_fonts.attrib[qn(key)]
        r_fonts.set(qn("w:ascii"), font_en)
        r_fonts.set(qn("w:hAnsi"), font_en)
        r_fonts.set(qn("w:cs"), font_en)
        r_fonts.set(qn("w:eastAsia"), font_zh)
        sz = ensure_child(r_pr, "w:sz")
        sz.set(qn("w:val"), str(int(size_pt * 2)))
        sz_cs = ensure_child(r_pr, "w:szCs")
        sz_cs.set(qn("w:val"), str(int(size_pt * 2)))
        lang = ensure_child(r_pr, "w:lang")
        lang.set(qn("w:val"), "en-US")
        lang.set(qn("w:eastAsia"), "zh-CN")
        p_pr_default = ensure_child(doc_defaults, "w:pPrDefault")
        p_pr = ensure_child(p_pr_default, "w:pPr")
        spacing = ensure_child(p_pr, "w:spacing")
        spacing.set(qn("w:before"), str(int(space_before_pt * 20)))
        spacing.set(qn("w:after"), str(int(space_after_pt * 20)))
        spacing.set(qn("w:line"), str(int(line_spacing * 240)))
        spacing.set(qn("w:lineRule"), "auto")

    def configure_list_style(style, font_zh: str, font_en: str, size_pt: float, line_spacing: float, left_indent_cm: float,
                             hanging_indent_cm: float, align: str) -> None:
        set_style_fonts(style, font_zh, font_en)
        style.font.size = Pt(size_pt)
        set_spacing(style, line_spacing, 0, 0, align=align)
        style.paragraph_format.left_indent = Cm(left_indent_cm)
        style.paragraph_format.first_line_indent = Cm(-hanging_indent_cm)

    def set_run_text(paragraph, text: str, font_zh: str, font_en: str, size_pt: float) -> None:
        paragraph.text = ""
        run = paragraph.add_run(text)
        run.font.name = font_en
        run.font.size = Pt(size_pt)
        r_pr = run._element.get_or_add_rPr()
        r_fonts = ensure_child(r_pr, "w:rFonts")
        r_fonts.set(qn("w:ascii"), font_en)
        r_fonts.set(qn("w:hAnsi"), font_en)
        r_fonts.set(qn("w:cs"), font_en)
        r_fonts.set(qn("w:eastAsia"), font_zh)

    document = Document()
    section = document.sections[0]
    page = profile["page"]
    margins = page.get("margins_cm")
    if margins:
        section.top_margin = Cm(float(margins["top"]))
        section.bottom_margin = Cm(float(margins["bottom"]))
        section.left_margin = Cm(float(margins["left"]))
        section.right_margin = Cm(float(margins["right"]))
    else:
        margin_cm = float(page["margin_cm"])
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)
    if page.get("header_distance_cm") is not None:
        section.header_distance = Cm(float(page["header_distance_cm"]))
    if page.get("footer_distance_cm") is not None:
        section.footer_distance = Cm(float(page["footer_distance_cm"]))

    body = profile["body"]
    configure_doc_defaults(
        document,
        body["font_zh"],
        body["font_en"],
        float(body["size_pt"]),
        float(body["line_spacing"]),
        float(body.get("space_before_pt", 0)),
        float(body["space_after_pt"]),
    )

    normal = document.styles["Normal"]
    configure_body_style(
        normal,
        body["font_zh"],
        body["font_en"],
        float(body["size_pt"]),
        float(body["line_spacing"]),
        float(body.get("space_before_pt", 0)),
        float(body["space_after_pt"]),
        float(body.get("first_line_indent_pt", 24)),
        body.get("align", "justify"),
    )

    for style_name, base_name in (("Body Text", "Normal"), ("First Paragraph", "Body Text"), ("Compact", "Body Text")):
        style = ensure_paragraph_style(document, style_name, base_name)
        configure_body_style(
            style,
            body["font_zh"],
            body["font_en"],
            float(body["size_pt"]),
            float(body["line_spacing"]),
            float(body.get("space_before_pt", 0)),
            0 if style_name == "Compact" else float(body["space_after_pt"]),
            0 if style_name == "Compact" else float(body.get("first_line_indent_pt", 24)),
            body.get("align", "justify"),
        )
        if style_name == "Compact":
            style.paragraph_format.first_line_indent = Pt(0)

    for style_name, heading_name in (
        ("heading1", "Heading 1"),
        ("heading2", "Heading 2"),
        ("heading3", "Heading 3"),
        ("heading4", "Heading 4"),
    ):
        if style_name not in profile["headings"]:
            continue
        style = ensure_paragraph_style(document, heading_name)
        style = document.styles[heading_name]
        config = profile["headings"][style_name]
        configure_heading_style(
            style,
            config["font_zh"],
            config["font_en"],
            float(config["size_pt"]),
            bool(config["bold"]),
            float(config["space_before_pt"]),
            float(config["space_after_pt"]),
            float(config.get("line_spacing", profile["body"]["line_spacing"])),
            config.get("align", "left"),
        )
        style.paragraph_format.first_line_indent = Pt(0)

    configure_title_style(document.styles["Title"], body["font_zh"], body["font_en"])

    captions = profile.get("captions", {})
    for style_name, base_name in (("Caption", "Body Text"), ("Table Caption", "Caption"), ("Image Caption", "Caption")):
        style = ensure_paragraph_style(document, style_name, base_name)
        configure_caption_style(
            style,
            captions.get("font_zh", body["font_zh"]),
            captions.get("font_en", body["font_en"]),
            float(captions.get("size_pt", body["size_pt"])),
            float(captions.get("line_spacing", body["line_spacing"])),
            float(captions.get("space_before_pt", 6)),
            float(captions.get("space_after_pt", 6)),
            captions.get("align", "center"),
        )

    lists = profile.get("lists", {})
    for style_name in ("List Paragraph", "List Bullet", "List Number"):
        style = ensure_paragraph_style(document, style_name, "Body Text")
        configure_list_style(
            style,
            body["font_zh"],
            body["font_en"],
            float(body["size_pt"]),
            float(body["line_spacing"]),
            float(lists.get("left_indent_cm", 0.84)),
            float(lists.get("hanging_indent_cm", 0.84)),
            body.get("align", "justify"),
        )

    header_footer = profile.get("header_footer", {})
    section.different_first_page_header_footer = bool(header_footer.get("different_first_page", False))
    document.settings.odd_and_even_pages_header_footer = bool(header_footer.get("different_odd_even", False))
    if header_footer.get("header_text"):
        header_para = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if header_footer.get("header_align", "center") == "center" else WD_ALIGN_PARAGRAPH.LEFT
        set_run_text(
            header_para,
            str(header_footer["header_text"]),
            str(header_footer.get("header_font_zh", body["font_zh"])),
            str(header_footer.get("header_font_en", body["font_en"])),
            float(header_footer.get("header_size_pt", 9)),
        )
    if header_footer.get("page_number", False):
        footer_para = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(fld_end)

    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate or patch a reference.docx file from a built-in template profile.")
    parser.add_argument("--template", required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--font")
    parser.add_argument("--font-size")
    parser.add_argument("--margin-cm")
    parser.add_argument("--dump-json", action="store_true")
    args = parser.parse_args()
    profile = load_profile(args.template)
    patched = patch_profile(profile, {"font": args.font or "", "font_size": args.font_size or "", "margin_cm": args.margin_cm or ""})
    if args.dump_json:
        print(json.dumps(patched, ensure_ascii=False, indent=2))
        return 0
    args.output.parent.mkdir(parents=True, exist_ok=True)
    generate_docx(patched, args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
